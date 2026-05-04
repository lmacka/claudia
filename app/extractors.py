"""
Per-format extractors for the library pipeline.

Each Extractor takes the original file path and returns:
- ExtractResult — the markdown to write to extracted.md, the extractor name,
  optional page_count, and any extra meta (chat-export participants etc).
- VerifyResult — heuristic and optional model-spot-check results that go into
  verification.json.
- DateDetection — best-effort original_date, with a source label.

The orchestrator (route handler) picks an extractor via ExtractorRegistry.pick,
then calls extract → verify → detect_date in sequence, emitting status
messages along the way via the supplied Emit callback.
"""

from __future__ import annotations

import datetime as _dt
import io
import logging
import mimetypes
import re
import shutil
import subprocess
import tempfile
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Protocol

import structlog

log = structlog.get_logger()


# ---------------------------------------------------------------------------
# Types
# ---------------------------------------------------------------------------


Emit = Callable[[str], None]
"""Status-line callback. The orchestrator pushes these to the SSE stream."""


@dataclass
class ExtractResult:
    extracted_md: str
    extractor: str
    page_count: int | None = None
    extra_meta: dict[str, Any] = field(default_factory=dict)


@dataclass
class VerifyResult:
    status: str  # "ok" | "warn" | "fail"
    checks: list[dict[str, Any]] = field(default_factory=list)


@dataclass
class DateDetection:
    date: _dt.date | None
    source: str  # one of library.DateSource literals, or "unknown"


def _noop_emit(_msg: str) -> None:
    """Default emit: discard. Tests use list-collecting emitters."""


# ---------------------------------------------------------------------------
# Extractor protocol
# ---------------------------------------------------------------------------


class Extractor(Protocol):
    kind: str

    def can_handle(self, path: Path, mime: str) -> bool: ...
    def extract(self, path: Path, emit: Emit = _noop_emit) -> ExtractResult: ...
    def verify(
        self, path: Path, result: ExtractResult, emit: Emit = _noop_emit
    ) -> VerifyResult: ...
    def detect_date(self, path: Path) -> DateDetection: ...


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------


class ExtractorRegistry:
    """
    Pick-by-mime, with a chat-export sniffer for text payloads. The order
    of extractors matters — chat detection runs before generic text.
    """

    def __init__(self, extractors: list[Extractor] | None = None) -> None:
        self._extractors: list[Extractor] = list(extractors or [])

    def register(self, extractor: Extractor) -> None:
        self._extractors.append(extractor)

    def pick(self, path: Path, mime: str | None = None) -> Extractor | None:
        if mime is None:
            mime, _ = mimetypes.guess_type(str(path))
            mime = mime or "application/octet-stream"
        for ex in self._extractors:
            try:
                if ex.can_handle(path, mime):
                    return ex
            except Exception as e:  # noqa: BLE001 — extractor sniffers must not crash registry
                logging.getLogger(__name__).warning(
                    "extractor sniffer failed: %s on %s — %s", ex.kind, path.name, e
                )
        return None

    def __iter__(self):
        return iter(self._extractors)


# ---------------------------------------------------------------------------
# TextExtractor
# ---------------------------------------------------------------------------


class TextExtractor:
    """
    Verbatim copy. Title heuristic = first non-empty line.

    The library entry's `title` is a caller concern (paste form gives the
    user a label, file uploads default to filename); this extractor only
    deals with content. The title heuristic helper is exposed for callers
    that want it.
    """

    kind = "text"

    def can_handle(self, path: Path, mime: str) -> bool:
        if mime.startswith("text/"):
            return True
        # Catch .md and .txt that mimetypes sometimes mis-guesses.
        return path.suffix.lower() in (".txt", ".md", ".markdown", ".log")

    def extract(self, path: Path, emit: Emit = _noop_emit) -> ExtractResult:
        emit("Reading text…")
        text = path.read_text(encoding="utf-8", errors="replace")
        return ExtractResult(extracted_md=text, extractor="text_verbatim")

    def verify(
        self, path: Path, result: ExtractResult, emit: Emit = _noop_emit
    ) -> VerifyResult:
        emit("Verifying byte count…")
        original_bytes = path.stat().st_size
        # Round-trip equality on UTF-8 char count is fuzzy because of decode
        # errors; use the byte count of the encoded extracted as the check.
        round_trip_bytes = len(result.extracted_md.encode("utf-8"))
        ok = abs(round_trip_bytes - original_bytes) <= max(8, original_bytes // 100)
        return VerifyResult(
            status="ok" if ok else "warn",
            checks=[
                {
                    "name": "byte_round_trip",
                    "ok": ok,
                    "detail": f"original={original_bytes}B, extracted={round_trip_bytes}B",
                }
            ],
        )

    def detect_date(self, path: Path) -> DateDetection:
        # Free-text pastes have no automatic date. Caller prompts the user.
        return DateDetection(date=None, source="unknown")

    @staticmethod
    def title_from_text(text: str, fallback: str = "Untitled note") -> str:
        for line in text.splitlines():
            stripped = line.strip()
            if stripped:
                return stripped[:80]
        return fallback


# ---------------------------------------------------------------------------
# ImageExtractor
# ---------------------------------------------------------------------------


VisionTranscribeFn = Callable[[bytes, str], str]
"""(image_bytes, mime_type) -> markdown transcript."""


VisionSpotCheckFn = Callable[[bytes, str, str], str]
"""(image_bytes, mime_type, extracted_md) -> 'ok' or list of issues."""


VISION_TRANSCRIBE_PROMPT = (
    "Transcribe all visible text in this image verbatim. "
    "Note any UI chrome, sender names, timestamps. "
    "Use markdown for headings and lists. "
    "If the image contains no text, respond with exactly: NO_TEXT_FOUND."
)

VISION_SPOTCHECK_PROMPT = (
    "Below is a transcription claimed to be from the attached image. "
    "Compare them. Respond with exactly 'ok' if the transcription is "
    "complete and accurate, or list specific issues (one per line). "
    "Do not include any other commentary.\n\n"
    "Transcription:\n{extracted}"
)


class ImageExtractor:
    """
    Vision OCR via injected callable (Sonnet in prod, mocked in tests).
    EXIF date detection via Pillow.
    """

    kind = "image"
    SUPPORTED_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".heic")
    SPOT_CHECK_DEFAULT = True

    def __init__(
        self,
        transcribe: VisionTranscribeFn | None = None,
        spot_check: VisionSpotCheckFn | None = None,
        spot_check_enabled: bool = SPOT_CHECK_DEFAULT,
    ) -> None:
        self._transcribe = transcribe
        self._spot_check = spot_check
        self._spot_check_enabled = spot_check_enabled

    def can_handle(self, path: Path, mime: str) -> bool:
        return mime.startswith("image/") or path.suffix.lower() in self.SUPPORTED_EXTS

    def extract(self, path: Path, emit: Emit = _noop_emit) -> ExtractResult:
        if self._transcribe is None:
            raise RuntimeError(
                "ImageExtractor requires a transcribe callable; pass one at construction"
            )
        emit("Reading image…")
        image_bytes = path.read_bytes()
        mime = self._mime_for(path)
        emit("Transcribing via vision…")
        text = self._transcribe(image_bytes, mime)
        if text.strip() == "NO_TEXT_FOUND":
            text = ""
        return ExtractResult(extracted_md=text, extractor="image_vision_ocr")

    def verify(
        self, path: Path, result: ExtractResult, emit: Emit = _noop_emit
    ) -> VerifyResult:
        checks: list[dict[str, Any]] = []
        chars = len(result.extracted_md)
        # Heuristic: image with no text is a warn, not a fail (memes / photos).
        no_text_ok = chars > 0
        checks.append(
            {
                "name": "transcript_nonempty",
                "ok": no_text_ok,
                "detail": f"extracted={chars} chars (warn if 0; some images have no text)",
            }
        )
        if not no_text_ok:
            return VerifyResult(status="warn", checks=checks)

        if self._spot_check_enabled and self._spot_check is not None:
            emit("Sanity checking…")
            try:
                image_bytes = path.read_bytes()
                mime = self._mime_for(path)
                judgement = self._spot_check(image_bytes, mime, result.extracted_md).strip()
                ok = judgement.lower() == "ok"
                checks.append(
                    {
                        "name": "vision_spot_check",
                        "ok": ok,
                        "detail": "ok" if ok else f"issues: {judgement[:200]}",
                    }
                )
                if not ok:
                    return VerifyResult(status="warn", checks=checks)
            except Exception as e:  # noqa: BLE001
                checks.append(
                    {
                        "name": "vision_spot_check",
                        "ok": False,
                        "detail": f"spot-check call failed: {e!s}",
                    }
                )
                return VerifyResult(status="warn", checks=checks)

        return VerifyResult(status="ok", checks=checks)

    def detect_date(self, path: Path) -> DateDetection:
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
        except ImportError:
            return DateDetection(date=None, source="unknown")
        try:
            with Image.open(path) as img:
                exif = img.getexif()
                if not exif:
                    return DateDetection(date=None, source="unknown")
                # Tag 36867 = DateTimeOriginal. Format: "YYYY:MM:DD HH:MM:SS"
                for tag_id, value in exif.items():
                    name = TAGS.get(tag_id, tag_id)
                    if name in ("DateTimeOriginal", "DateTime", "DateTimeDigitized"):
                        if isinstance(value, str):
                            parsed = self._parse_exif_date(value)
                            if parsed:
                                return DateDetection(date=parsed, source="exif")
        except Exception as e:  # noqa: BLE001 — date detection must not crash extraction
            log.warning("image.exif_read_failed", path=str(path), error=str(e))
        return DateDetection(date=None, source="unknown")

    @staticmethod
    def _parse_exif_date(s: str) -> _dt.date | None:
        try:
            return _dt.datetime.strptime(s.split(" ")[0], "%Y:%m:%d").date()
        except (ValueError, IndexError):
            return None

    @staticmethod
    def _mime_for(path: Path) -> str:
        guessed, _ = mimetypes.guess_type(str(path))
        if guessed:
            return guessed
        return f"image/{path.suffix.lstrip('.').lower() or 'png'}"


# ---------------------------------------------------------------------------
# PdfExtractor
# ---------------------------------------------------------------------------


# Used by the pypdf↔pdfplumber cross-check to normalise whitespace before
# substring comparison, so trivial spacing differences don't trigger warnings.
_WHITESPACE_RUN = re.compile(r"\s+")

# Date patterns swept on PDF first page when metadata is missing.
_PDF_DATE_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b"), "iso"),
    (re.compile(r"\b(\d{1,2})\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+(\d{4})\b", re.I), "dmy"),
    (re.compile(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+(\d{1,2}),?\s+(\d{4})\b", re.I), "mdy"),
    (re.compile(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b"), "slash_dmy"),
]
_MONTH_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}


class PdfExtractor:
    """
    Primary path: pypdf over all pages, no page cap. Per-page exceptions
    are recorded inline and don't abort. If extracted-chars-per-page falls
    below SCAN_THRESHOLD, fall back to vision OCR via pdf2image rasterize.
    """

    kind = "pdf"
    SCAN_THRESHOLD = 100  # chars/page; below = treat as scanned PDF
    EMIT_EVERY_PAGES = 5  # emit status every N pages for big PDFs
    SPOT_CHECK_DEFAULT = True

    def __init__(
        self,
        transcribe: VisionTranscribeFn | None = None,
        spot_check: VisionSpotCheckFn | None = None,
        spot_check_enabled: bool = SPOT_CHECK_DEFAULT,
    ) -> None:
        self._transcribe = transcribe
        self._spot_check = spot_check
        self._spot_check_enabled = spot_check_enabled

    def can_handle(self, path: Path, mime: str) -> bool:
        return mime == "application/pdf" or path.suffix.lower() == ".pdf"

    def extract(self, path: Path, emit: Emit = _noop_emit) -> ExtractResult:
        from pypdf import PdfReader

        emit("Opening PDF…")
        reader = PdfReader(str(path))
        n_pages = len(reader.pages)
        if n_pages == 0:
            return ExtractResult(extracted_md="", extractor="pdf_pypdf", page_count=0)

        pages_md: list[str] = []
        for i, page in enumerate(reader.pages, start=1):
            if i == 1 or i == n_pages or i % self.EMIT_EVERY_PAGES == 0:
                emit(f"Extracting page {i}/{n_pages}…")
            try:
                text = page.extract_text() or ""
            except Exception as e:  # noqa: BLE001
                pages_md.append(f"## Page {i}\n\n[page {i}: extraction failed: {e!s}]\n")
                continue
            pages_md.append(f"## Page {i}\n\n{text.strip()}\n" if text.strip() else f"## Page {i}\n\n[page {i}: no text extracted]\n")

        extracted_md = "\n".join(pages_md)
        chars_per_page = len(extracted_md) / n_pages if n_pages > 0 else 0

        if chars_per_page < self.SCAN_THRESHOLD and n_pages > 0:
            emit("Sparse text detected, falling back to vision OCR…")
            return self._extract_via_vision(path, n_pages, emit)

        return ExtractResult(
            extracted_md=extracted_md,
            extractor="pdf_pypdf",
            page_count=n_pages,
        )

    def _extract_via_vision(self, path: Path, n_pages: int, emit: Emit) -> ExtractResult:
        if self._transcribe is None:
            # Best-effort: return whatever pypdf got (probably nothing) but
            # tag the extractor honestly so the verifier flags it.
            return ExtractResult(
                extracted_md=f"[scanned PDF, no vision transcribe available; {n_pages} pages]",
                extractor="pdf_vision_ocr_unavailable",
                page_count=n_pages,
            )

        from pdf2image import convert_from_path

        emit("Rasterising pages for OCR…")
        images = convert_from_path(str(path), dpi=200)
        pages_md: list[str] = []
        for i, img in enumerate(images, start=1):
            if i == 1 or i == n_pages or i % self.EMIT_EVERY_PAGES == 0:
                emit(f"OCR page {i}/{n_pages}…")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            try:
                text = self._transcribe(buf.getvalue(), "image/png")
            except Exception as e:  # noqa: BLE001
                text = f"[page {i}: vision OCR failed: {e!s}]"
            pages_md.append(f"## Page {i}\n\n{text.strip()}\n")
        return ExtractResult(
            extracted_md="\n".join(pages_md),
            extractor="pdf_vision_ocr",
            page_count=n_pages,
        )

    def verify(
        self, path: Path, result: ExtractResult, emit: Emit = _noop_emit
    ) -> VerifyResult:
        checks: list[dict[str, Any]] = []
        n_pages = result.page_count or 0
        chars = len(result.extracted_md)

        # Heuristic 1: chars per page
        if n_pages > 0:
            cpp = chars / n_pages
            cpp_ok = cpp >= 50
            checks.append(
                {
                    "name": "chars_per_page",
                    "ok": cpp_ok,
                    "detail": f"{cpp:.1f} chars/page (warn if <50)",
                }
            )
        else:
            cpp_ok = True

        # Heuristic 2: pdfplumber cross-check on text-extracted PDFs
        cross_ok = True
        if result.extractor == "pdf_pypdf" and n_pages > 0:
            emit("Cross-checking with pdfplumber…")
            try:
                cross_ok = self._cross_check_pdfplumber(path, result.extracted_md, n_pages)
                checks.append(
                    {
                        "name": "pdfplumber_cross_check",
                        "ok": cross_ok,
                        "detail": "first/last/middle pages overlap with pypdf output"
                        if cross_ok
                        else "pdfplumber found text pypdf missed; possible extraction drift",
                    }
                )
            except Exception as e:  # noqa: BLE001
                checks.append(
                    {
                        "name": "pdfplumber_cross_check",
                        "ok": False,
                        "detail": f"cross-check failed: {e!s}",
                    }
                )
                cross_ok = False

        # Spot-check (vision-OCR'd PDFs only): page 1 image vs page 1 transcript.
        if (
            result.extractor == "pdf_vision_ocr"
            and self._spot_check_enabled
            and self._spot_check is not None
            and n_pages > 0
        ):
            emit("Sanity checking page 1…")
            try:
                page1_md = result.extracted_md.split("\n## Page 2\n", 1)[0]
                from pdf2image import convert_from_path

                imgs = convert_from_path(str(path), dpi=200, first_page=1, last_page=1)
                if imgs:
                    buf = io.BytesIO()
                    imgs[0].save(buf, format="PNG")
                    judgement = self._spot_check(buf.getvalue(), "image/png", page1_md).strip()
                    spot_ok = judgement.lower() == "ok"
                    checks.append(
                        {
                            "name": "vision_spot_check_page_1",
                            "ok": spot_ok,
                            "detail": "ok" if spot_ok else f"issues: {judgement[:200]}",
                        }
                    )
                    if not spot_ok:
                        cross_ok = False
            except Exception as e:  # noqa: BLE001
                checks.append(
                    {
                        "name": "vision_spot_check_page_1",
                        "ok": False,
                        "detail": f"spot-check failed: {e!s}",
                    }
                )

        status = "ok" if (cpp_ok and cross_ok) else "warn"
        return VerifyResult(status=status, checks=checks)

    @staticmethod
    def _cross_check_pdfplumber(path: Path, pypdf_text: str, n_pages: int) -> bool:
        import pdfplumber

        # Whitespace handling differs between pypdf and pdfplumber (pypdf often
        # adds trailing spaces, pdfplumber preserves \n where pypdf may insert
        # a space). A literal substring match flags those whitespace deltas as
        # "missing text" even though the content is identical. Normalize both
        # sides — collapse any whitespace run to a single space — before the
        # 30-char snippet comparison so the check fires only on genuine
        # extraction drift.
        norm_pypdf = _WHITESPACE_RUN.sub(" ", pypdf_text).strip()
        sample_indices = {0, n_pages - 1, max(0, n_pages // 2)}
        with pdfplumber.open(str(path)) as pdf:
            for idx in sorted(sample_indices):
                if idx >= len(pdf.pages):
                    continue
                plumber_text = pdf.pages[idx].extract_text() or ""
                norm_plumber = _WHITESPACE_RUN.sub(" ", plumber_text).strip()
                snippet = norm_plumber[:30]
                if snippet and snippet not in norm_pypdf:
                    return False
        return True

    def detect_date(self, path: Path) -> DateDetection:
        from pypdf import PdfReader

        try:
            reader = PdfReader(str(path))
            meta = reader.metadata
            if meta:
                created = getattr(meta, "creation_date", None)
                if isinstance(created, _dt.datetime):
                    return DateDetection(date=created.date(), source="pdf_metadata")
                raw = meta.get("/CreationDate") if hasattr(meta, "get") else None
                if isinstance(raw, str):
                    parsed = self._parse_pdf_creation_date(raw)
                    if parsed:
                        return DateDetection(date=parsed, source="pdf_metadata")
            # Regex sweep on first page text.
            if reader.pages:
                text = reader.pages[0].extract_text() or ""
                for pat, kind in _PDF_DATE_PATTERNS:
                    m = pat.search(text)
                    if m:
                        parsed = self._date_from_match(m, kind)
                        if parsed:
                            return DateDetection(date=parsed, source="pdf_text_pattern")
        except Exception as e:  # noqa: BLE001
            log.warning("pdf.date_detect_failed", path=str(path), error=str(e))
        return DateDetection(date=None, source="unknown")

    @staticmethod
    def _parse_pdf_creation_date(raw: str) -> _dt.date | None:
        # Format: "D:YYYYMMDDHHmmSSOHH'mm'" — only need the date part.
        if raw.startswith("D:"):
            raw = raw[2:]
        if len(raw) < 8:
            return None
        try:
            return _dt.date(int(raw[0:4]), int(raw[4:6]), int(raw[6:8]))
        except (ValueError, IndexError):
            return None

    @staticmethod
    def _date_from_match(m: re.Match, kind: str) -> _dt.date | None:
        try:
            if kind == "iso":
                return _dt.date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
            if kind == "dmy":
                return _dt.date(int(m.group(3)), _MONTH_MAP[m.group(2).lower()[:3]], int(m.group(1)))
            if kind == "mdy":
                return _dt.date(int(m.group(3)), _MONTH_MAP[m.group(1).lower()[:3]], int(m.group(2)))
            if kind == "slash_dmy":
                # Ambiguous DMY/MDY; claudia is AU, so DMY first.
                d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
                if d > 12 and mo <= 12:
                    return _dt.date(y, mo, d)
                return _dt.date(y, mo, d)  # default DMY interpretation
        except (ValueError, KeyError, IndexError):
            return None
        return None


# ---------------------------------------------------------------------------
# DocxExtractor
# ---------------------------------------------------------------------------


class DocxExtractor:
    """python-docx walks paragraphs (heading-level aware), tables, headers/footers."""

    kind = "docx"

    def can_handle(self, path: Path, mime: str) -> bool:
        if path.suffix.lower() == ".docx":
            return True
        return mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/octet-stream",
        ) and path.suffix.lower() == ".docx"

    def extract(self, path: Path, emit: Emit = _noop_emit) -> ExtractResult:
        from docx import Document

        emit("Opening DOCX…")
        doc = Document(str(path))

        lines: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name.startswith("heading 1"):
                lines.append(f"# {text}")
            elif style_name.startswith("heading 2"):
                lines.append(f"## {text}")
            elif style_name.startswith("heading 3"):
                lines.append(f"### {text}")
            elif style_name.startswith("heading 4"):
                lines.append(f"#### {text}")
            elif style_name.startswith("list"):
                lines.append(f"- {text}")
            else:
                lines.append(text)

        # Tables → markdown
        for table in doc.tables:
            lines.append("")
            for row_i, row in enumerate(table.rows):
                cells = [c.text.strip().replace("|", "\\|") for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if row_i == 0:
                    lines.append("|" + "|".join("---" for _ in cells) + "|")
            lines.append("")

        # Headers/footers
        h_f_lines: list[str] = []
        for section in doc.sections:
            for header in (section.header,):
                hpieces = [p.text.strip() for p in header.paragraphs if p.text.strip()]
                if hpieces:
                    h_f_lines.append("**Header:** " + " · ".join(hpieces))
            for footer in (section.footer,):
                fpieces = [p.text.strip() for p in footer.paragraphs if p.text.strip()]
                if fpieces:
                    h_f_lines.append("**Footer:** " + " · ".join(fpieces))
        if h_f_lines:
            lines.append("")
            lines.append("## Headers/footers")
            lines.append("")
            lines.extend(h_f_lines)

        # Inline images: just a count
        inline_images = sum(1 for shape in doc.inline_shapes)
        if inline_images:
            lines.append("")
            lines.append(f"## Embedded images ({inline_images})")
            lines.append("")
            lines.append("(extraction skipped per spec — out of scope to save embedded images as separate library docs)")

        paragraph_count = sum(1 for p in doc.paragraphs)
        return ExtractResult(
            extracted_md="\n".join(lines).strip() + "\n",
            extractor="docx_python_docx",
            page_count=None,
            extra_meta={"paragraph_count": paragraph_count, "inline_image_count": inline_images},
        )

    def verify(
        self, path: Path, result: ExtractResult, emit: Emit = _noop_emit
    ) -> VerifyResult:
        from docx import Document

        para_count = result.extra_meta.get("paragraph_count", 0)
        checks: list[dict[str, Any]] = [
            {"name": "paragraph_count", "ok": para_count > 0, "detail": f"{para_count} paragraphs"}
        ]

        emit("Round-trip checking…")
        try:
            Document(str(path))  # re-open after write — a smoke test
            checks.append({"name": "round_trip_open", "ok": True, "detail": "re-opens cleanly"})
        except Exception as e:  # noqa: BLE001
            checks.append({"name": "round_trip_open", "ok": False, "detail": str(e)})
            return VerifyResult(status="warn", checks=checks)

        return VerifyResult(status="ok" if para_count > 0 else "warn", checks=checks)

    def detect_date(self, path: Path) -> DateDetection:
        try:
            from docx import Document

            doc = Document(str(path))
            created = doc.core_properties.created
            if isinstance(created, _dt.datetime):
                return DateDetection(date=created.date(), source="docx_core_props")
        except Exception as e:  # noqa: BLE001
            log.warning("docx.date_detect_failed", path=str(path), error=str(e))
        return DateDetection(date=None, source="unknown")


# ---------------------------------------------------------------------------
# DocExtractor — legacy .doc via LibreOffice → DOCX
# ---------------------------------------------------------------------------


class DocExtractor:
    """
    Legacy .doc files. Converts to DOCX in a tempdir via LibreOffice headless,
    then delegates to DocxExtractor. Original .doc preserved as original.doc.
    """

    kind = "doc"
    LIBREOFFICE_BIN = "libreoffice"
    CONVERSION_TIMEOUT_S = 60

    def __init__(self, docx_extractor: DocxExtractor | None = None) -> None:
        self._docx = docx_extractor or DocxExtractor()

    def can_handle(self, path: Path, mime: str) -> bool:
        return path.suffix.lower() == ".doc" or mime == "application/msword"

    def _convert(self, path: Path) -> Path:
        if shutil.which(self.LIBREOFFICE_BIN) is None:
            raise RuntimeError(
                f"{self.LIBREOFFICE_BIN!r} not found in PATH. Install libreoffice-core to handle .doc files."
            )
        tmp = Path(tempfile.mkdtemp(prefix="claudia-doc-"))
        try:
            result = subprocess.run(
                [
                    self.LIBREOFFICE_BIN,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(tmp),
                    str(path),
                ],
                capture_output=True,
                text=True,
                timeout=self.CONVERSION_TIMEOUT_S,
                check=False,
            )
            if result.returncode != 0:
                raise RuntimeError(
                    f"libreoffice convert exited {result.returncode}: {result.stderr.strip()[:300]}"
                )
            converted = tmp / (path.stem + ".docx")
            if not converted.exists():
                raise RuntimeError(
                    f"libreoffice did not produce expected output {converted}; stdout={result.stdout.strip()[:300]}"
                )
            return converted
        except subprocess.TimeoutExpired as e:
            shutil.rmtree(tmp, ignore_errors=True)
            raise RuntimeError(f"libreoffice conversion timed out after {self.CONVERSION_TIMEOUT_S}s") from e

    def extract(self, path: Path, emit: Emit = _noop_emit) -> ExtractResult:
        emit("Converting .doc → .docx via LibreOffice…")
        converted = self._convert(path)
        try:
            result = self._docx.extract(converted, emit)
            return ExtractResult(
                extracted_md=result.extracted_md,
                extractor="doc_libreoffice_docx",
                page_count=result.page_count,
                extra_meta=result.extra_meta,
            )
        finally:
            shutil.rmtree(converted.parent, ignore_errors=True)

    def verify(
        self, path: Path, result: ExtractResult, emit: Emit = _noop_emit
    ) -> VerifyResult:
        # We delegated extraction; verify is a paragraph-count smoke check
        # on the result (the converted docx is gone). This is enough to flag
        # a conversion that succeeded but produced empty output.
        para_count = result.extra_meta.get("paragraph_count", 0)
        ok = para_count > 0
        return VerifyResult(
            status="ok" if ok else "warn",
            checks=[{"name": "paragraph_count", "ok": ok, "detail": f"{para_count} paragraphs"}],
        )

    def detect_date(self, path: Path) -> DateDetection:
        try:
            converted = self._convert(path)
            try:
                return self._docx.detect_date(converted)
            finally:
                shutil.rmtree(converted.parent, ignore_errors=True)
        except Exception as e:  # noqa: BLE001
            log.warning("doc.date_detect_failed", path=str(path), error=str(e))
            return DateDetection(date=None, source="unknown")


# ---------------------------------------------------------------------------
# ChatExportExtractor — WhatsApp / iMessage / Telegram-shaped pastes
# ---------------------------------------------------------------------------


# WhatsApp: "[12/03/24, 14:32:01] Liam: hello" or
#           "[12/03/2024, 2:32 PM] Liam: hello"
# Square brackets, date dd/mm/yy or yyyy, comma optional, time HH:MM(:SS)( AM/PM)?
_WHATSAPP_LINE = re.compile(
    r"^\[(\d{1,2})/(\d{1,2})/(\d{2,4}),?\s+(\d{1,2}):(\d{2})(?::(\d{2}))?\s*(AM|PM)?\]\s+([^:]+):\s*(.*)$",
    re.IGNORECASE,
)
_SYSTEM_MARKERS = (
    "Messages and calls are end-to-end encrypted",
    "<This message was edited>",
    "image omitted",
    "video omitted",
    "audio omitted",
    "Media omitted",
    "<Media omitted>",
    "GIF omitted",
    "sticker omitted",
    "Missed voice call",
    "Missed video call",
)


@dataclass
class _ChatMessage:
    timestamp: _dt.datetime
    sender: str
    body: str


class ChatExportExtractor:
    """
    Detects WhatsApp-shaped pastes / .txt files. Sniff: >50% of first 20
    non-empty lines match the WhatsApp regex.
    """

    kind = "chat_export"
    SNIFF_LINES = 20
    SNIFF_THRESHOLD = 0.5

    def can_handle(self, path: Path, mime: str) -> bool:
        # Cheap: only consider text-shaped payloads
        if not (mime.startswith("text/") or path.suffix.lower() in (".txt", ".log", "")):
            return False
        try:
            with path.open("r", encoding="utf-8", errors="replace") as fh:
                head = "".join([line for _i, line in zip(range(2000), fh, strict=False)])
        except OSError:
            return False
        non_empty = [ln for ln in head.splitlines() if ln.strip()][: self.SNIFF_LINES]
        if len(non_empty) < 5:
            return False
        matches = sum(1 for ln in non_empty if _WHATSAPP_LINE.match(ln))
        return matches / len(non_empty) >= self.SNIFF_THRESHOLD

    def extract(self, path: Path, emit: Emit = _noop_emit) -> ExtractResult:
        emit("Parsing chat export…")
        original_lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
        messages, system_count = self._parse(original_lines)
        if not messages:
            return ExtractResult(
                extracted_md="(no chat messages parsed)",
                extractor="chat_whatsapp",
                extra_meta={"message_count": 0, "participants": [], "system_message_count": system_count},
            )

        emit("Rendering markdown…")
        participants = sorted({m.sender for m in messages})
        first = messages[0].timestamp
        last = messages[-1].timestamp
        header = (
            f"## Chat export — {', '.join(participants)} — "
            f"{first.date().isoformat()} → {last.date().isoformat()}\n"
        )
        body_lines = [header, ""]
        for msg in messages:
            ts = msg.timestamp.strftime("%Y-%m-%d %H:%M")
            body_lines.append(f"**{ts} — {msg.sender}:** {msg.body}")
        if system_count:
            body_lines.append("")
            body_lines.append(f"*(stripped {system_count} system messages)*")

        return ExtractResult(
            extracted_md="\n".join(body_lines) + "\n",
            extractor="chat_whatsapp",
            extra_meta={
                "message_count": len(messages),
                "participants": participants,
                "system_message_count": system_count,
                "first_message_date": first.date().isoformat(),
                "last_message_date": last.date().isoformat(),
            },
        )

    @staticmethod
    def _parse(lines: list[str]) -> tuple[list[_ChatMessage], int]:
        messages: list[_ChatMessage] = []
        system_count = 0
        current: _ChatMessage | None = None
        for raw in lines:
            line = raw.rstrip("\r")
            m = _WHATSAPP_LINE.match(line)
            if m:
                # Flush previous before starting new
                if current is not None:
                    messages.append(current)
                ts = ChatExportExtractor._build_ts(m)
                if ts is None:
                    current = None
                    continue
                sender = m.group(8).strip()
                body = m.group(9).strip()
                if any(marker in body for marker in _SYSTEM_MARKERS) or sender == "":
                    system_count += 1
                    current = None
                    continue
                current = _ChatMessage(timestamp=ts, sender=sender, body=body)
            elif current is not None and line.strip():
                # Continuation of previous message
                current.body += "\n" + line
            elif not line.strip() and current is not None:
                messages.append(current)
                current = None
        if current is not None:
            messages.append(current)
        return messages, system_count

    @staticmethod
    def _build_ts(m: re.Match) -> _dt.datetime | None:
        try:
            d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            if y < 100:
                y += 2000
            hour = int(m.group(4))
            minute = int(m.group(5))
            second = int(m.group(6) or 0)
            ampm = m.group(7)
            if ampm:
                if ampm.upper() == "PM" and hour < 12:
                    hour += 12
                elif ampm.upper() == "AM" and hour == 12:
                    hour = 0
            # WhatsApp dates are usually DD/MM/YY in AU; if d>12, that's a hint;
            # default to DMY which matches AU export conventions.
            if d > 12 and mo <= 12:
                pass  # already DMY
            return _dt.datetime(y, mo, d, hour, minute, second)
        except (ValueError, KeyError):
            return None

    def verify(
        self, path: Path, result: ExtractResult, emit: Emit = _noop_emit
    ) -> VerifyResult:
        original_lines = [ln for ln in path.read_text(encoding="utf-8", errors="replace").splitlines() if ln.strip()]
        msg_count = result.extra_meta.get("message_count", 0)
        sys_count = result.extra_meta.get("system_message_count", 0)
        target = max(1, len(original_lines) - sys_count)
        ratio = msg_count / target
        ratio_ok = ratio >= 0.9
        checks = [
            {
                "name": "parse_coverage",
                "ok": ratio_ok,
                "detail": f"{msg_count}/{target} non-system lines parsed ({ratio:.2%})",
            },
        ]

        # Sender presence: every detected participant should appear in markdown.
        participants = result.extra_meta.get("participants", [])
        senders_in_md = all(p in result.extracted_md for p in participants)
        checks.append(
            {
                "name": "all_senders_in_output",
                "ok": senders_in_md,
                "detail": f"{len(participants)} unique senders detected",
            }
        )
        ok = ratio_ok and senders_in_md
        return VerifyResult(status="ok" if ok else "warn", checks=checks)

    def detect_date(self, path: Path) -> DateDetection:
        # Need to parse to know the first message timestamp.
        try:
            lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
            messages, _ = self._parse(lines)
            if messages:
                return DateDetection(date=messages[0].timestamp.date(), source="chat_first_message")
        except Exception as e:  # noqa: BLE001
            log.warning("chat.date_detect_failed", path=str(path), error=str(e))
        return DateDetection(date=None, source="unknown")


# ---------------------------------------------------------------------------
# Build a default registry. Order matters: chat sniffer before generic text.
# ---------------------------------------------------------------------------


def build_registry(
    transcribe: VisionTranscribeFn | None = None,
    spot_check: VisionSpotCheckFn | None = None,
) -> ExtractorRegistry:
    return ExtractorRegistry(
        [
            PdfExtractor(transcribe=transcribe, spot_check=spot_check),
            DocxExtractor(),
            DocExtractor(),
            ImageExtractor(transcribe=transcribe, spot_check=spot_check),
            ChatExportExtractor(),  # before TextExtractor — sniffs text payloads
            TextExtractor(),  # last — generic fallback
        ]
    )


# ---------------------------------------------------------------------------
# Helpers shared with route handlers
# ---------------------------------------------------------------------------


def make_vision_callables(
    claude_client: Any,
    model: str = "claude-sonnet-4-6",
) -> tuple[VisionTranscribeFn, VisionSpotCheckFn]:
    """
    Build the transcribe + spot_check callables that wrap a ClaudeClient.
    Imported here (and not at module top) to avoid a hard claude.py
    dependency at import time — useful for tests that don't need vision.
    """

    def transcribe(image_bytes: bytes, mime: str) -> str:
        import base64

        b64 = base64.b64encode(image_bytes).decode("ascii")
        reply = claude_client.single_turn(
            model=model,
            user_content=[
                {"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}},
                {"type": "text", "text": VISION_TRANSCRIBE_PROMPT},
            ],
            max_tokens=4096,
        )
        return reply.text

    def spot_check(image_bytes: bytes, mime: str, extracted_md: str) -> str:
        import base64

        b64 = base64.b64encode(image_bytes).decode("ascii")
        prompt = VISION_SPOTCHECK_PROMPT.format(extracted=extracted_md[:4000])
        reply = claude_client.single_turn(
            model=model,
            user_content=[
                {"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}},
                {"type": "text", "text": prompt},
            ],
            max_tokens=512,
        )
        return reply.text

    # Discourage the linter from removing the io import (used by callers
    # constructing in-memory bytes for testing).
    _ = io.BytesIO
    return transcribe, spot_check
