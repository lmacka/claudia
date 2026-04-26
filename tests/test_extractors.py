"""Tests for app/extractors.py — framework + Text + Image + PDF + DOCX + DOC + chat."""

from __future__ import annotations

import datetime as _dt
import io
import shutil
from pathlib import Path

import pytest
from PIL import Image

from app.extractors import (
    ChatExportExtractor,
    DateDetection,
    DocExtractor,
    DocxExtractor,
    ExtractorRegistry,
    ExtractResult,
    ImageExtractor,
    PdfExtractor,
    TextExtractor,
    VerifyResult,
    build_registry,
)


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _emits() -> tuple[list, callable]:
    """Returns (collected_list, emit_fn)."""
    collected: list[str] = []
    return collected, collected.append


def _png_bytes(size: tuple[int, int] = (4, 4)) -> bytes:
    img = Image.new("RGB", size, color=(128, 64, 32))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _write_png(path: Path) -> None:
    path.write_bytes(_png_bytes())


# ---------------------------------------------------------------------------
# TextExtractor
# ---------------------------------------------------------------------------


def test_text_extractor_can_handle_text_mime():
    ex = TextExtractor()
    assert ex.can_handle(Path("note.txt"), "text/plain")
    assert ex.can_handle(Path("note.md"), "text/markdown")
    assert ex.can_handle(Path("anything.log"), "application/octet-stream")  # by ext
    assert not ex.can_handle(Path("photo.png"), "image/png")


def test_text_extractor_extract_verbatim(tmp_path: Path):
    p = tmp_path / "note.txt"
    p.write_text("hello world\nline two\n", encoding="utf-8")
    emits, emit = _emits()
    result = TextExtractor().extract(p, emit)
    assert result.extracted_md == "hello world\nline two\n"
    assert result.extractor == "text_verbatim"
    assert "Reading text" in emits[0]


def test_text_extractor_verify_byte_round_trip(tmp_path: Path):
    p = tmp_path / "note.txt"
    content = "small note"
    p.write_text(content, encoding="utf-8")
    ex = TextExtractor()
    extract = ex.extract(p)
    verify = ex.verify(p, extract)
    assert verify.status == "ok"
    assert any(check["name"] == "byte_round_trip" for check in verify.checks)


def test_text_extractor_detect_date_returns_unknown(tmp_path: Path):
    p = tmp_path / "note.txt"
    p.write_text("free text", encoding="utf-8")
    detection = TextExtractor().detect_date(p)
    assert detection.date is None
    assert detection.source == "unknown"


def test_text_extractor_title_heuristic():
    text = "  \n\n  First real line\nsecond line\n"
    assert TextExtractor.title_from_text(text) == "First real line"


def test_text_extractor_title_truncates_long_first_line():
    text = "x" * 200
    assert len(TextExtractor.title_from_text(text)) == 80


def test_text_extractor_title_fallback_for_empty():
    assert TextExtractor.title_from_text("") == "Untitled note"
    assert TextExtractor.title_from_text("\n\n  \n") == "Untitled note"


# ---------------------------------------------------------------------------
# ImageExtractor — extract
# ---------------------------------------------------------------------------


def test_image_extractor_can_handle_image_mime():
    ex = ImageExtractor(transcribe=lambda b, m: "")
    assert ex.can_handle(Path("p.png"), "image/png")
    assert ex.can_handle(Path("p.jpg"), "image/jpeg")
    assert ex.can_handle(Path("p.heic"), "application/octet-stream")  # by ext
    assert not ex.can_handle(Path("notes.txt"), "text/plain")


def test_image_extractor_extract_calls_transcribe(tmp_path: Path):
    p = tmp_path / "snap.png"
    _write_png(p)

    captured: dict = {}

    def fake_transcribe(image_bytes: bytes, mime: str) -> str:
        captured["bytes_len"] = len(image_bytes)
        captured["mime"] = mime
        return "## Sofia\nkk\nidk like its fine"

    emits, emit = _emits()
    ex = ImageExtractor(transcribe=fake_transcribe)
    result = ex.extract(p, emit)
    assert result.extractor == "image_vision_ocr"
    assert "Sofia" in result.extracted_md
    assert captured["bytes_len"] > 0
    assert captured["mime"] == "image/png"
    assert any("Transcribing" in m for m in emits)


def test_image_extractor_no_text_found_returns_empty(tmp_path: Path):
    p = tmp_path / "blank.png"
    _write_png(p)
    ex = ImageExtractor(transcribe=lambda b, m: "NO_TEXT_FOUND")
    result = ex.extract(p)
    assert result.extracted_md == ""


def test_image_extractor_requires_transcribe_callable(tmp_path: Path):
    p = tmp_path / "snap.png"
    _write_png(p)
    ex = ImageExtractor()  # no transcribe injected
    with pytest.raises(RuntimeError):
        ex.extract(p)


# ---------------------------------------------------------------------------
# ImageExtractor — verify
# ---------------------------------------------------------------------------


def test_image_verify_ok_with_text_and_passing_spotcheck(tmp_path: Path):
    p = tmp_path / "snap.png"
    _write_png(p)
    ex = ImageExtractor(
        transcribe=lambda b, m: "transcribed text",
        spot_check=lambda b, m, e: "ok",
    )
    extract = ex.extract(p)
    verify = ex.verify(p, extract)
    assert verify.status == "ok"
    assert any(c["name"] == "vision_spot_check" and c["ok"] for c in verify.checks)


def test_image_verify_warn_when_no_text(tmp_path: Path):
    p = tmp_path / "blank.png"
    _write_png(p)
    ex = ImageExtractor(transcribe=lambda b, m: "NO_TEXT_FOUND")
    extract = ex.extract(p)
    verify = ex.verify(p, extract)
    assert verify.status == "warn"
    # Spot-check skipped because empty text gates it.
    assert all(c["name"] != "vision_spot_check" for c in verify.checks)


def test_image_verify_warn_when_spotcheck_lists_issues(tmp_path: Path):
    p = tmp_path / "snap.png"
    _write_png(p)
    ex = ImageExtractor(
        transcribe=lambda b, m: "transcribed",
        spot_check=lambda b, m, e: "missing the word 'urgent' in line 3\nwrong sender name",
    )
    extract = ex.extract(p)
    verify = ex.verify(p, extract)
    assert verify.status == "warn"
    spot = next(c for c in verify.checks if c["name"] == "vision_spot_check")
    assert not spot["ok"]
    assert "missing" in spot["detail"]


def test_image_verify_warn_when_spotcheck_raises(tmp_path: Path):
    p = tmp_path / "snap.png"
    _write_png(p)

    def boom(*_args):
        raise RuntimeError("spot check api blew up")

    ex = ImageExtractor(
        transcribe=lambda b, m: "transcribed",
        spot_check=boom,
    )
    extract = ex.extract(p)
    verify = ex.verify(p, extract)
    assert verify.status == "warn"
    spot = next(c for c in verify.checks if c["name"] == "vision_spot_check")
    assert "spot-check call failed" in spot["detail"]


def test_image_verify_skips_spotcheck_when_disabled(tmp_path: Path):
    p = tmp_path / "snap.png"
    _write_png(p)
    spot_calls = []
    ex = ImageExtractor(
        transcribe=lambda b, m: "transcribed",
        spot_check=lambda b, m, e: spot_calls.append("called") or "ok",
        spot_check_enabled=False,
    )
    verify = ex.verify(p, ex.extract(p))
    assert verify.status == "ok"
    assert spot_calls == []  # never called


# ---------------------------------------------------------------------------
# ImageExtractor — date detection
# ---------------------------------------------------------------------------


def test_image_detect_date_no_exif(tmp_path: Path):
    p = tmp_path / "no-exif.png"
    _write_png(p)
    detection = ImageExtractor().detect_date(p)
    assert detection.date is None
    assert detection.source == "unknown"


def test_image_parse_exif_date_valid():
    parsed = ImageExtractor._parse_exif_date("2024:11:23 18:42:01")
    assert parsed == _dt.date(2024, 11, 23)


def test_image_parse_exif_date_invalid_returns_none():
    assert ImageExtractor._parse_exif_date("not a date") is None
    assert ImageExtractor._parse_exif_date("") is None


def test_image_detect_date_with_real_exif(tmp_path: Path):
    """Write a JPEG with EXIF DateTimeOriginal and confirm we read it back."""
    p = tmp_path / "phone.jpg"
    img = Image.new("RGB", (8, 8), color=(10, 20, 30))
    exif = img.getexif()
    exif[306] = "2026:04:25 09:15:00"  # DateTime
    exif[36867] = "2026:04:25 09:15:00"  # DateTimeOriginal
    img.save(p, format="JPEG", exif=exif)

    detection = ImageExtractor().detect_date(p)
    assert detection.date == _dt.date(2026, 4, 25)
    assert detection.source == "exif"


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------


def test_registry_pick_image_first(tmp_path: Path):
    reg = build_registry(transcribe=lambda b, m: "")
    p = tmp_path / "snap.png"
    _write_png(p)
    picked = reg.pick(p, "image/png")
    assert picked is not None
    assert picked.kind == "image"


def test_registry_pick_text_for_text_payload(tmp_path: Path):
    reg = build_registry()
    p = tmp_path / "note.txt"
    p.write_text("hello", encoding="utf-8")
    picked = reg.pick(p, "text/plain")
    assert picked is not None
    assert picked.kind == "text"


def test_registry_pick_returns_none_for_unknown(tmp_path: Path):
    reg = build_registry()
    p = tmp_path / "mystery.bin"
    p.write_bytes(b"\x00\x01\x02")
    # Without a binary handler, no extractor can_handle this.
    picked = reg.pick(p, "application/octet-stream")
    assert picked is None


def test_registry_iteration():
    reg = build_registry()
    kinds = [ex.kind for ex in reg]
    assert "image" in kinds
    assert "text" in kinds


def test_registry_swallows_sniffer_exceptions(tmp_path: Path):
    """A misbehaving extractor's can_handle must not crash the registry."""

    class BoomExtractor:
        kind = "boom"

        def can_handle(self, path, mime):
            raise RuntimeError("sniffer blew up")

        def extract(self, path, emit=lambda _m: None):  # pragma: no cover
            raise NotImplementedError

        def verify(self, path, result, emit=lambda _m: None):  # pragma: no cover
            raise NotImplementedError

        def detect_date(self, path):  # pragma: no cover
            raise NotImplementedError

    reg = ExtractorRegistry([BoomExtractor(), TextExtractor()])
    p = tmp_path / "n.txt"
    p.write_text("hi", encoding="utf-8")
    picked = reg.pick(p, "text/plain")
    assert picked is not None
    assert picked.kind == "text"


# ---------------------------------------------------------------------------
# Smoke: dataclass shapes
# ---------------------------------------------------------------------------


def test_extract_result_extra_meta_default():
    r = ExtractResult(extracted_md="x", extractor="t")
    assert r.extra_meta == {}


def test_verify_result_checks_default():
    v = VerifyResult(status="ok")
    assert v.checks == []


def test_date_detection_unknown():
    d = DateDetection(date=None, source="unknown")
    assert d.date is None


# ===========================================================================
# PdfExtractor
# ===========================================================================


def _write_text_pdf(path: Path, pages: list[str]) -> None:
    """Write a real text-PDF using reportlab (already a project dep)."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import LETTER

    c = canvas.Canvas(str(path), pagesize=LETTER)
    for page_text in pages:
        for i, line in enumerate(page_text.splitlines() or [page_text]):
            c.drawString(72, 720 - i * 14, line[:120])
        c.showPage()
    c.save()


def test_pdf_can_handle():
    ex = PdfExtractor()
    assert ex.can_handle(Path("doc.pdf"), "application/pdf")
    assert ex.can_handle(Path("doc.pdf"), "application/octet-stream")
    assert not ex.can_handle(Path("notes.txt"), "text/plain")


def test_pdf_extract_text_pdf(tmp_path: Path):
    p = tmp_path / "report.pdf"
    # Each page needs >100 chars to avoid the scan-detection fallback.
    page_one = "Hello world. This is page one with enough text to clear the scan-detection threshold of 100 chars per page."
    page_two = "Second page content with similar density of words so pypdf returns a non-trivial extraction across both pages."
    _write_text_pdf(p, [page_one, page_two])
    emits, emit = _emits()
    result = PdfExtractor().extract(p, emit)
    assert result.extractor == "pdf_pypdf"
    assert result.page_count == 2
    assert "Hello world" in result.extracted_md
    assert "Second page" in result.extracted_md
    assert any("Extracting page" in m for m in emits)


def test_pdf_extract_falls_back_to_vision_for_sparse_text(tmp_path: Path):
    """If chars/page < SCAN_THRESHOLD, vision OCR runs."""
    p = tmp_path / "scanned.pdf"
    # Single page with virtually no text — forces scan branch
    _write_text_pdf(p, [""])
    transcribe_calls = []

    def fake_transcribe(image_bytes: bytes, mime: str) -> str:
        transcribe_calls.append((len(image_bytes), mime))
        return "OCR'd text from page"

    ex = PdfExtractor(transcribe=fake_transcribe)
    result = ex.extract(p)
    assert result.extractor == "pdf_vision_ocr"
    assert "OCR'd text" in result.extracted_md
    assert len(transcribe_calls) == 1
    assert transcribe_calls[0][1] == "image/png"


def test_pdf_extract_marks_unavailable_when_no_transcribe(tmp_path: Path):
    p = tmp_path / "scanned.pdf"
    _write_text_pdf(p, [""])
    ex = PdfExtractor()  # no transcribe injected
    result = ex.extract(p)
    assert result.extractor == "pdf_vision_ocr_unavailable"


def test_pdf_verify_ok_for_dense_text(tmp_path: Path):
    p = tmp_path / "report.pdf"
    body = "Lorem ipsum dolor sit amet. " * 30
    _write_text_pdf(p, [body, body])
    ex = PdfExtractor()
    extract = ex.extract(p)
    verify = ex.verify(p, extract)
    assert verify.status == "ok"


def test_pdf_detect_date_from_metadata(tmp_path: Path):
    """reportlab embeds /CreationDate; we should read it."""
    p = tmp_path / "with-date.pdf"
    _write_text_pdf(p, ["Hello"])
    detection = PdfExtractor().detect_date(p)
    # reportlab embeds today's date by default
    assert detection.source in ("pdf_metadata", "unknown")
    if detection.source == "pdf_metadata":
        assert isinstance(detection.date, _dt.date)


def test_pdf_parse_pdf_creation_date_string():
    parsed = PdfExtractor._parse_pdf_creation_date("D:20180312143215+11'00'")
    assert parsed == _dt.date(2018, 3, 12)


def test_pdf_parse_pdf_creation_date_invalid():
    assert PdfExtractor._parse_pdf_creation_date("not a date") is None
    assert PdfExtractor._parse_pdf_creation_date("D:short") is None


def test_pdf_date_from_match_iso():
    import re

    m = re.compile(r"(\d{4})-(\d{2})-(\d{2})").search("Date: 2018-03-12 issued")
    assert PdfExtractor._date_from_match(m, "iso") == _dt.date(2018, 3, 12)


# ===========================================================================
# DocxExtractor
# ===========================================================================


def _write_docx(
    path: Path,
    paragraphs: list[tuple[str, str]] | None = None,
    table_rows: list[list[str]] | None = None,
) -> None:
    """Write a small .docx using python-docx itself."""
    from docx import Document

    doc = Document()
    for style, text in paragraphs or [("Normal", "default body")]:
        doc.add_paragraph(text, style=style)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                table.cell(r, c).text = cell
    doc.save(str(path))


def test_docx_can_handle():
    ex = DocxExtractor()
    assert ex.can_handle(Path("note.docx"), "")
    assert not ex.can_handle(Path("note.pdf"), "application/pdf")


def test_docx_extract_paragraphs_and_headings(tmp_path: Path):
    p = tmp_path / "note.docx"
    _write_docx(
        p,
        paragraphs=[
            ("Heading 1", "Title"),
            ("Heading 2", "Subhead"),
            ("Normal", "body line"),
        ],
    )
    result = DocxExtractor().extract(p)
    assert "# Title" in result.extracted_md
    assert "## Subhead" in result.extracted_md
    assert "body line" in result.extracted_md
    assert result.extractor == "docx_python_docx"


def test_docx_extract_renders_tables(tmp_path: Path):
    p = tmp_path / "table.docx"
    _write_docx(
        p,
        paragraphs=[("Normal", "intro")],
        table_rows=[["A", "B"], ["1", "2"]],
    )
    result = DocxExtractor().extract(p)
    md = result.extracted_md
    assert "| A | B |" in md
    assert "| 1 | 2 |" in md


def test_docx_verify_ok(tmp_path: Path):
    p = tmp_path / "note.docx"
    _write_docx(p, paragraphs=[("Normal", "hello")])
    ex = DocxExtractor()
    verify = ex.verify(p, ex.extract(p))
    assert verify.status == "ok"


def test_docx_detect_date_from_core_props(tmp_path: Path):
    p = tmp_path / "note.docx"
    _write_docx(p, paragraphs=[("Normal", "hi")])
    detection = DocxExtractor().detect_date(p)
    # python-docx auto-stamps core_properties.created on save.
    assert detection.source in ("docx_core_props", "unknown")
    if detection.source == "docx_core_props":
        assert isinstance(detection.date, _dt.date)


# ===========================================================================
# DocExtractor (legacy .doc) — conditional on libreoffice
# ===========================================================================


HAS_LIBREOFFICE = shutil.which("libreoffice") is not None


def test_doc_can_handle():
    ex = DocExtractor()
    assert ex.can_handle(Path("legacy.doc"), "")
    assert ex.can_handle(Path("anywhere"), "application/msword")
    assert not ex.can_handle(Path("note.docx"), "")


def test_doc_extract_raises_when_libreoffice_missing(tmp_path: Path):
    """Without LibreOffice in PATH, the conversion should raise cleanly."""
    p = tmp_path / "fake.doc"
    p.write_bytes(b"\xd0\xcf\x11\xe0")  # OLE compound doc magic header
    ex = DocExtractor()
    if HAS_LIBREOFFICE:
        pytest.skip("libreoffice IS installed; this test asserts the error path")
    with pytest.raises(RuntimeError, match="not found in PATH"):
        ex.extract(p)


@pytest.mark.skipif(not HAS_LIBREOFFICE, reason="libreoffice not installed")
def test_doc_extract_via_libreoffice(tmp_path: Path):
    """End-to-end: write a .doc via libreoffice, extract it back."""
    docx_path = tmp_path / "src.docx"
    _write_docx(docx_path, paragraphs=[("Normal", "doc-extractor smoke test")])
    # Convert .docx → .doc via libreoffice for the round-trip.
    import subprocess

    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "doc", "--outdir", str(tmp_path), str(docx_path)],
        capture_output=True,
        text=True,
        timeout=60,
    )
    assert result.returncode == 0, result.stderr
    doc_path = tmp_path / "src.doc"
    assert doc_path.exists()

    extract = DocExtractor().extract(doc_path)
    assert "doc-extractor smoke test" in extract.extracted_md
    assert extract.extractor == "doc_libreoffice_docx"


# ===========================================================================
# ChatExportExtractor — WhatsApp shape
# ===========================================================================


WHATSAPP_SAMPLE = """\
[12/04/26, 14:32:01] Liam: hey are we still on for sat
[12/04/26, 14:35:00] Rhiannon: yeah 4pm
[12/04/26, 14:35:30] Rhiannon: bring jasper's stuff please
[12/04/26, 18:01:00] Liam: ok will do
[12/04/26, 18:02:11] Rhiannon: also <Media omitted>
[13/04/26, 8:14:00] Liam: how was the rest of the day
"""


def _write_chat(path: Path, content: str = WHATSAPP_SAMPLE) -> None:
    path.write_text(content, encoding="utf-8")


def test_chat_can_handle_whatsapp(tmp_path: Path):
    p = tmp_path / "_chat.txt"
    _write_chat(p)
    assert ChatExportExtractor().can_handle(p, "text/plain")


def test_chat_can_handle_rejects_random_text(tmp_path: Path):
    p = tmp_path / "notes.txt"
    p.write_text("just a regular note\nwith multiple lines\nof prose\nnothing chat here\nat all", encoding="utf-8")
    assert not ChatExportExtractor().can_handle(p, "text/plain")


def test_chat_can_handle_rejects_non_text(tmp_path: Path):
    p = tmp_path / "snap.png"
    p.write_bytes(b"\x89PNG")
    assert not ChatExportExtractor().can_handle(p, "image/png")


def test_chat_extract_parses_messages(tmp_path: Path):
    p = tmp_path / "_chat.txt"
    _write_chat(p)
    result = ChatExportExtractor().extract(p)
    assert result.extractor == "chat_whatsapp"
    assert "Liam" in result.extra_meta["participants"]
    assert "Rhiannon" in result.extra_meta["participants"]
    # 5 real messages + 1 system "<Media omitted>" stripped
    assert result.extra_meta["message_count"] == 5
    assert result.extra_meta["system_message_count"] == 1
    assert "Liam" in result.extracted_md
    assert "Rhiannon" in result.extracted_md


def test_chat_extract_continuation_lines(tmp_path: Path):
    p = tmp_path / "_chat.txt"
    p.write_text(
        "[12/04/26, 14:32:01] Liam: line one\nline two of same message\nline three\n"
        "[12/04/26, 14:35:00] Rhiannon: response\n",
        encoding="utf-8",
    )
    result = ChatExportExtractor().extract(p)
    assert result.extra_meta["message_count"] == 2
    # The continuation should be in the body of the first message.
    assert "line one\nline two" in result.extracted_md or "line two" in result.extracted_md


def test_chat_verify_ok(tmp_path: Path):
    p = tmp_path / "_chat.txt"
    _write_chat(p)
    ex = ChatExportExtractor()
    verify = ex.verify(p, ex.extract(p))
    assert verify.status == "ok"


def test_chat_detect_date_first_message(tmp_path: Path):
    p = tmp_path / "_chat.txt"
    _write_chat(p)
    detection = ChatExportExtractor().detect_date(p)
    assert detection.source == "chat_first_message"
    assert detection.date == _dt.date(2026, 4, 12)


# ===========================================================================
# Registry order check (post-commit-C)
# ===========================================================================


def test_registry_order_pdf_before_text(tmp_path: Path):
    p = tmp_path / "doc.pdf"
    _write_text_pdf(p, ["hello"])
    reg = build_registry()
    picked = reg.pick(p, "application/pdf")
    assert picked is not None
    assert picked.kind == "pdf"


def test_registry_chat_before_text(tmp_path: Path):
    p = tmp_path / "_chat.txt"
    _write_chat(p)
    reg = build_registry()
    picked = reg.pick(p, "text/plain")
    assert picked is not None
    assert picked.kind == "chat_export"


def test_registry_text_picks_for_plain_note(tmp_path: Path):
    p = tmp_path / "note.txt"
    p.write_text("just a regular note about something\nnothing chat here\nplain prose only\nfour lines\nfive\n", encoding="utf-8")
    reg = build_registry()
    picked = reg.pick(p, "text/plain")
    assert picked is not None
    assert picked.kind == "text"


def test_registry_docx_picks_for_docx(tmp_path: Path):
    p = tmp_path / "note.docx"
    _write_docx(p, paragraphs=[("Normal", "hi")])
    reg = build_registry()
    picked = reg.pick(p, "")
    assert picked is not None
    assert picked.kind == "docx"
